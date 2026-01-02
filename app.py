from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, RGBColor, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from io import BytesIO
import socket
from contextlib import contextmanager
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError


from staticmap import StaticMap, CircleMarker
from PIL import Image
from datetime import datetime  # Manejo de fechas

app = Flask(__name__)

# --- Health check (Render / monitoring) ---

from flask import jsonify

@app.get("/health")
def health():
    return jsonify(ok=True)

@app.get("/")
def root():
    return jsonify(service="coe-word-service", ok=True)






# ============================================================
# FUNCIONES DE FORMATO COMUNES
# ============================================================

def set_cell_bg(cell, color_hex: str):
    """Aplica color de fondo a una celda (hex sin #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_section_title(doc: Document, text: str):
    """Crea faja amarilla con título azul (estilo MIDIS)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, "FEE599")

    p = cell.paragraphs[0]
    p.alignment = 0  # izquierda
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def configurar_cabeceras(doc: Document):
    """Inserta la cabecera COE en TODAS las páginas (RP/RC)."""
    for section in doc.sections:
        # Misma cabecera en primera y siguientes páginas
        section.different_first_page_header_footer = False

        # Asegura que el contenido no se monte con una cabecera alta (ajuste conservador)
        try:
            if section.top_margin < Cm(3.2):
                section.top_margin = Cm(3.2)
        except Exception:
            pass

        # Ancho exacto del área útil (página - márgenes)
        try:
            header_width = section.page_width - section.left_margin - section.right_margin
        except Exception:
            header_width = Inches(6.5)

        def _set_header(hdr):
            try:
                hdr.is_linked_to_previous = False
            except Exception:
                pass

            # deja un solo párrafo en cabecera
            if hdr.paragraphs:
                p = hdr.paragraphs[0]
                for extra in hdr.paragraphs[1:]:
                    try:
                        extra._element.getparent().remove(extra._element)
                    except Exception:
                        pass
                p.text = ""
            else:
                p = hdr.add_paragraph()

            p.alignment = 1  # centrado
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

            run = p.add_run()
            # intenta PNG, luego JPG
            try:
                run.add_picture("cabecera_coe_1.png", width=header_width)
            except Exception:
                try:
                    run.add_picture("cabecera_coe_1.jpg", width=header_width)
                except Exception:
                    pass

        _set_header(section.header)
        # por compatibilidad, también setea first_page_header (aunque esté desactivado)
        try:
            _set_header(section.first_page_header)
        except Exception:
            pass


def set_paragraph_single_spacing(p):
    """Espaciado anterior 0, posterior 0, interlineado línea única."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE



@contextmanager
def _socket_timeout(seconds: float):
    """Aplica un timeout por defecto a sockets (útil para teselas OSM)."""
    old = socket.getdefaulttimeout()
    socket.setdefaulttimeout(seconds)
    try:
        yield
    finally:
        socket.setdefaulttimeout(old)


def _render_static_map(lat_f: float, lon_f: float, zoom: int):
    """Renderiza un mapa estático (puede fallar/colgar si no hay teselas)."""
    m = StaticMap(800, 600, url_template="https://tile.openstreetmap.org/{z}/{x}/{y}.png")
    marker = CircleMarker((lon_f, lat_f), "#d50000", 12)
    m.add_marker(marker)
    return m.render(zoom=zoom)

def ensure_paragraph_runs(p):
    """Asegura que el párrafo tenga al menos un run."""
    if not p.runs:
        p.add_run("")
    return p.runs


def obtener_dt_elaboracion(data: dict):
    """
    Obtiene datetime desde el payload. Soporta varias llaves (RP/RC) y formatos típicos.
    Preferencia: fechaHoraRC / fechaElaboracionRC / fechaElaboracion / fechaHora / fechaRegistro
    """
    for key in ("fechaHoraRC", "fechaElaboracionRC", "fechaElaboracion", "fechaHora", "fechaRegistro"):
        fh = str(data.get(key, "") or "").strip()
        if not fh:
            continue
        # normaliza ISO con Z
        if fh.endswith("Z"):
            fh = fh[:-1]
        try:
            return datetime.fromisoformat(fh)
        except Exception:
            # intento simple: "YYYY-MM-DD HH:MM"
            try:
                return datetime.strptime(fh, "%Y-%m-%d %H:%M")
            except Exception:
                pass
    return None

def formatear_fecha_ddmmyyyy(fecha_iso: str):
    """Convierte 'YYYY-MM-DD' → 'DD-MM-YYYY'. Si falla, devuelve la original."""
    try:
        d = datetime.strptime(fecha_iso, "%Y-%m-%d")
        return d.strftime("%d-%m-%Y")
    except:
        return fecha_iso


def insertar_tabla_ubicacion_y_mapa(doc: Document, data: dict):
    """
    Inserta:
    - Tabla: Departamento / Provincia / Distrito
    - Mapa generado localmente (staticmap), tamaño 12 cm × 8 cm
    """
    dep = str(data.get("departamento", "") or "")
    prov = str(data.get("provincia", "") or "")
    dist = str(data.get("distrito", "") or "")
    lat_raw = (data.get("latitud") or "").strip()
    lon_raw = (data.get("longitud") or "").strip()
    peligro = (data.get("peligro") or "").lower()

    # Tabla de ubicación
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Departamento", "Provincia", "Distrito"

    for c in hdr:
        runs = ensure_paragraph_runs(c.paragraphs[0])
        for r in runs:
            r.bold = True
            r.font.name = "Calibri"

    vals = table.rows[1].cells
    vals[0].text, vals[1].text, vals[2].text = dep, prov, dist
    for c in vals:
        runs = ensure_paragraph_runs(c.paragraphs[0])
        for r in runs:
            r.font.name = "Calibri"

    # Título mapa
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(3)
    p_title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p_title.add_run("Mapa de ubicación")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)

    # Validación de coordenadas
    if not lat_raw or not lon_raw:
        p = doc.add_paragraph("Sin coordenadas registradas (no se puede generar el mapa).")
        runs = ensure_paragraph_runs(p)
        for r in runs:
            r.font.name = "Calibri"
        return

    lat_str = lat_raw.replace(",", ".")
    lon_str = lon_raw.replace(",", ".")

    try:
        lat_f = float(lat_str)
        lon_f = float(lon_str)
    except Exception:
        p = doc.add_paragraph(
            f"No se pudo interpretar las coordenadas: latitud='{lat_raw}', longitud='{lon_raw}'."
        )
        runs = ensure_paragraph_runs(p)
        for r in runs:
            r.font.name = "Calibri"
        return

    # Zoom según peligro (aprox.)
    zoom = 13
    if "sismo" in peligro:
        zoom = 10
    elif "huaic" in peligro:
        zoom = 14
    elif "inund" in peligro or "lluvia" in peligro:
        zoom = 13
    elif "incendios urbanos" in peligro or "incendio urbano" in peligro:
        zoom = 15
    elif "incendios forestales" in peligro or "incendio forestal" in peligro:
        zoom = 13

    try:
        # Render de mapa: proteger contra demoras (teselas OSM / red)
        img = None
        # 1) timeout de sockets (por si la librería se queda esperando)
        with _socket_timeout(2.8):
            # 2) además, ejecutar en hilo y cortar a los ~3s
            with ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(_render_static_map, lat_f, lon_f, zoom)
                try:
                    img = fut.result(timeout=3.0)
                except FuturesTimeoutError:
                    img = None

        if img is None:
            raise TimeoutError("timeout mapa")

        stream = BytesIO()
        img.save(stream, format="PNG")
        stream.seek(0)
        p_map = doc.add_paragraph()
        run_map = p_map.add_run()
        run_map.add_picture(stream, width=Cm(12), height=Cm(8))

    except Exception:
        p = doc.add_paragraph(
            "No se pudo generar el mapa estático (error al obtener las teselas de mapa). "
            f"Coordenadas: {lat_raw}, {lon_raw}."
        )
        runs = ensure_paragraph_runs(p)
        for r in runs:
            r.font.name = "Calibri"


# ============================================================
# API: GENERAR WORD RP
# ============================================================

@app.route("/api/generar-word-rp", methods=["POST"])
def generar_word_rp():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Sin datos"}), 400

    doc = Document()
    configurar_cabeceras(doc)

    # ---------------------------------------------
    # ENCABEZADO PERSONALIZADO RP
    # ---------------------------------------------
    peligro = (data.get("peligro") or "").upper()
    distrito = (data.get("distrito") or "").upper()
    departamento = (data.get("departamento") or "").upper()
    num_global = str(data.get("numeroGlobal", data.get("numeroReporte", "")))

    # 1) TÍTULO: PELIGRO EN EL DISTRITO X – Y
    titulo = f"{peligro} EN EL DISTRITO {distrito} – {departamento}".strip()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1  # centrado
    set_paragraph_single_spacing(p_titulo)
    run_t = p_titulo.add_run(titulo)
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # 2) FECHA DE ELABORACIÓN
    dt = obtener_dt_elaboracion(data)
    if dt:
        fecha_texto = dt.strftime("%d/%m/%Y %H:%M")
        fecha_archivo = dt.strftime("%d%m%Y")
    else:
        fecha_texto = ""
        fecha_archivo = ""

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = 1
    set_paragraph_single_spacing(p_fecha)
    r_fecha = p_fecha.add_run(f"Fecha de elaboración : {fecha_texto}")
    r_fecha.bold = True
    r_fecha.font.name = "Calibri"
    r_fecha.font.size = Pt(9)
    r_fecha.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # 3) CÓDIGO DE EMERGENCIA
    codigo = str(data.get("codigo", "") or "")
    p_cod = doc.add_paragraph()
    p_cod.alignment = 1
    set_paragraph_single_spacing(p_cod)
    r_cod = p_cod.add_run(f"Código de Emergencia: {codigo}")
    r_cod.bold = True
    r_cod.font.name = "Calibri"
    r_cod.font.size = Pt(9)
    # Color rojo D50000
    r_cod.font.color.rgb = RGBColor(0xD5, 0x00, 0x00)

    # 4) TEXTO REPORTE PRELIMINAR + N° GLOBAL
    p_rp = doc.add_paragraph()
    p_rp.alignment = 1
    set_paragraph_single_spacing(p_rp)
    r_rp = p_rp.add_run(f"REPORTE PRELIMINAR DE EMERGENCIA (RP) N° {num_global}")
    r_rp.bold = True
    r_rp.font.name = "Calibri"
    r_rp.font.size = Pt(9)
    r_rp.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ---------------------------------------------
    # CONTENIDO
    # ---------------------------------------------

    # HECHOS (en lugar de PELIGRO, primera sección)
    add_section_title(doc, "Hechos")
    p = doc.add_paragraph(str(data.get("hechos", "")))
    for r in ensure_paragraph_runs(p):
        r.font.name = "Calibri"

    # Ubicación + Mapa
    add_section_title(doc, "Ubicación")
    insertar_tabla_ubicacion_y_mapa(doc, data)

    # Daños MIDIS
    add_section_title(doc, "Daños en el sector Desarrollo e Inclusión Social")
    danios = data.get("daniosMIDIS", {})
    if danios:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Programa"
        hdr[1].text = "Usuarios afectados"
        hdr[2].text = "Servicios afectados"
        hdr[3].text = "Usuarios afectados por servicios"
        hdr[4].text = "Usuarios fallecidos"
        hdr[5].text = "Módulo afectado"

        for prog, vals in danios.items():
            row = table.add_row().cells
            row[0].text = prog
            row[1].text = str(vals.get("usuariosAfectados", 0))
            row[2].text = str(vals.get("serviciosAfectados", 0))
            row[3].text = str(vals.get("usuariosPorServicios", 0))
            row[4].text = str(vals.get("usuariosFallecidos", 0))
            row[5].text = str(vals.get("moduloAfectado", 0))
    else:
        p = doc.add_paragraph("Sin información registrada.")
        for r in ensure_paragraph_runs(p):
            r.font.name = "Calibri"

    # Otros sectores
    add_section_title(doc, "Daños en otros sectores")
    p = doc.add_paragraph(str(data.get("daniosOtros", "")))
    for r in ensure_paragraph_runs(p):
        r.font.name = "Calibri"

    # Acciones preliminares
    add_section_title(doc, "Acciones del Sector Desarrollo e Inclusión Social (Preliminar)")
    acciones = data.get("accionesPreliminar", [])
    if acciones:
        for i, acc in enumerate(acciones, start=1):
            if isinstance(acc, dict):
                fecha_acc = acc.get("fecha", "")
                desc = acc.get("descripcion", "")
                fecha_fmt = formatear_fecha_ddmmyyyy(fecha_acc)
                linea = f"{i}. [{fecha_fmt}] {desc}" if fecha_acc else f"{i}. {desc}"
            else:
                linea = f"{i}. {acc}"
            p = doc.add_paragraph(linea)
            for r in ensure_paragraph_runs(p):
                r.font.name = "Calibri"
    else:
        p = doc.add_paragraph("Sin acciones preliminares registradas.")
        for r in ensure_paragraph_runs(p):
            r.font.name = "Calibri"

    # Responsables
    add_section_title(doc, "Responsables")
    p = doc.add_paragraph()
    p.add_run("Elaborado por: ").bold = True
    p.add_run(str(data.get("elaboradoPor", "")))

    p = doc.add_paragraph()
    p.add_run("Aprobado por: ").bold = True
    p.add_run(str(data.get("aprobadoPor", "")))

    # ---------------------------------------------
    # NOMBRE DE ARCHIVO: NGlobal_Codigo_Peligro_Distrito_Departamento_Fecha
    # ---------------------------------------------
    numero_global_str = num_global or ""
    codigo_for_name = codigo.replace("-", "_")
    peligro_name = (data.get("peligro") or "").replace(" ", "")
    distrito_name = (data.get("distrito") or "").replace(" ", "")
    dep_name = (data.get("departamento") or "").replace(" ", "")
    fecha_compact = fecha_archivo or ""

    fname = f"{numero_global_str}_{codigo_for_name}_{peligro_name}_{distrito_name}_{dep_name}_{fecha_compact}.docx"

    # Exportar
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=fname,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============================================================
# API: GENERAR WORD RC
# ============================================================

@app.route("/api/generar-word-rc", methods=["POST"])
def generar_word_rc():
    data = request.get_json()
    if not data:
        return jsonify({"error": "Sin datos"}), 400

    doc = Document()
    configurar_cabeceras(doc)

    # ---------------------------------------------
    # ENCABEZADO PERSONALIZADO RC
    # ---------------------------------------------
    peligro = (data.get("peligro") or "").upper()
    distrito = (data.get("distrito") or "").upper()
    departamento = (data.get("departamento") or "").upper()
    num_global = str(data.get("numeroGlobal", data.get("numeroReporteRC", "")))

    # Título
    titulo = f"{peligro} EN EL DISTRITO {distrito} – {departamento}".strip()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1
    set_paragraph_single_spacing(p_titulo)
    run_t = p_titulo.add_run(titulo)
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(12)
    run_t.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Fecha de elaboración
    dt = obtener_dt_elaboracion(data)
    if dt:
        fecha_texto = dt.strftime("%d/%m/%Y %H:%M")
        fecha_archivo = dt.strftime("%d%m%Y")
    else:
        fecha_texto = ""
        fecha_archivo = ""

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = 1
    set_paragraph_single_spacing(p_fecha)
    r_fecha = p_fecha.add_run(f"Fecha de elaboración : {fecha_texto}")
    r_fecha.bold = True
    r_fecha.font.name = "Calibri"
    r_fecha.font.size = Pt(9)
    r_fecha.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Código
    codigo = str(data.get("codigo", "") or "")
    p_cod = doc.add_paragraph()
    p_cod.alignment = 1
    set_paragraph_single_spacing(p_cod)
    r_cod = p_cod.add_run(f"Código de Emergencia: {codigo}")
    r_cod.bold = True
    r_cod.font.name = "Calibri"
    r_cod.font.size = Pt(9)
    # Rojo D50000
    r_cod.font.color.rgb = RGBColor(0xD5, 0x00, 0x00)

    # Texto RC + N° Global
    p_rc = doc.add_paragraph()
    p_rc.alignment = 1
    set_paragraph_single_spacing(p_rc)
    r_rc = p_rc.add_run(f"REPORTE COMPLEMENTARIO DE EMERGENCIA (RC) N° {num_global}")
    r_rc.bold = True
    r_rc.font.name = "Calibri"
    r_rc.font.size = Pt(9)
    r_rc.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # ---------------------------------------------
    # CONTENIDO
    # ---------------------------------------------
    # HECHOS (primera sección, en lugar de “Peligro”)
    add_section_title(doc, "Hechos")
    p = doc.add_paragraph(str(data.get("hechos", "")))
    for r in ensure_paragraph_runs(p):
        r.font.name = "Calibri"

    # Ubicación + mapa
    add_section_title(doc, "Ubicación")
    insertar_tabla_ubicacion_y_mapa(doc, data)

    # Daños MIDIS RC
    add_section_title(doc, "Daños en el sector Desarrollo e Inclusión Social")
    danios = data.get("daniosMIDIS", {})
    if danios:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Programa"
        hdr[1].text = "Usuarios afectados"
        hdr[2].text = "Servicios afectados"
        hdr[3].text = "Usuarios afectados por servicios"
        hdr[4].text = "Usuarios fallecidos"
        hdr[5].text = "Módulo afectado"

        for prog, vals in danios.items():
            row = table.add_row().cells
            row[0].text = prog
            row[1].text = str(vals.get("usuariosAfectados", 0))
            row[2].text = str(vals.get("serviciosAfectados", 0))
            row[3].text = str(vals.get("usuariosPorServicios", 0))
            row[4].text = str(vals.get("usuariosFallecidos", 0))
            row[5].text = str(vals.get("moduloAfectado", 0))
    else:
        p = doc.add_paragraph("Sin información registrada.")
        for r in ensure_paragraph_runs(p):
            r.font.name = "Calibri"

    # Otros sectores
    add_section_title(doc, "Daños en otros sectores")
    p = doc.add_paragraph(str(data.get("daniosOtros", "")))
    for r in ensure_paragraph_runs(p):
        r.font.name = "Calibri"

    # Acciones RP (preliminares)
    add_section_title(doc, "Acciones del Sector Desarrollo e Inclusión Social (Preliminar)")
    accionesRP = data.get("accionesPreliminar", [])
    if accionesRP:
        for i, acc in enumerate(accionesRP, start=1):
            if isinstance(acc, dict):
                fecha_acc = acc.get("fecha", "")
                desc = acc.get("descripcion", "")
                fecha_fmt = formatear_fecha_ddmmyyyy(fecha_acc)
                linea = f"{i}. [{fecha_fmt}] {desc}" if fecha_acc else f"{i}. {desc}"
            else:
                linea = f"{i}. {acc}"
            p = doc.add_paragraph(linea)
            for r in ensure_paragraph_runs(p):
                r.font.name = "Calibri"
    else:
        p = doc.add_paragraph("Sin acciones preliminares registradas.")
        for r in ensure_paragraph_runs(p):
            r.font.name = "Calibri"

    # Acciones RC
    add_section_title(doc, "Acciones del Sector Desarrollo e Inclusión Social (RC)")
    accionesRC = data.get("accionesRC", [])
    if accionesRC:
        for i, acc in enumerate(accionesRC, start=1):
            fecha_acc = acc.get("fecha", "")
            desc = acc.get("descripcion", "")
            fecha_fmt = formatear_fecha_ddmmyyyy(fecha_acc)
            linea = f"{i}. [{fecha_fmt}] {desc}" if fecha_acc else f"{i}. {desc}"
            p = doc.add_paragraph(linea)
            for r in ensure_paragraph_runs(p):
                r.font.name = "Calibri"
    else:
        p = doc.add_paragraph("Sin acciones complementarias registradas.")
        for r in ensure_paragraph_runs(p):
            r.font.name = "Calibri"

    # Responsables
    add_section_title(doc, "Responsables")
    p = doc.add_paragraph()
    p.add_run("Elaborado por: ").bold = True
    p.add_run(str(data.get("elaboradoPor", "")))

    p = doc.add_paragraph()
    p.add_run("Aprobado por: ").bold = True
    p.add_run(str(data.get("aprobadoPor", "")))

    # ---------------------------------------------
    # NOMBRE ARCHIVO RC: NGlobal_Codigo_Peligro_Distrito_Departamento_Fecha
    # ---------------------------------------------
    numero_global_str = num_global or ""
    codigo_for_name = codigo.replace("-", "_")
    peligro_name = (data.get("peligro") or "").replace(" ", "")
    distrito_name = (data.get("distrito") or "").replace(" ", "")
    dep_name = (data.get("departamento") or "").replace(" ", "")
    fecha_compact = fecha_archivo or ""

    fname = f"{numero_global_str}_{codigo_for_name}_{peligro_name}_{distrito_name}_{dep_name}_{fecha_compact}.docx"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name=fname,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============================================================
# EJECUCIÓN
# ============================================================

if __name__ == "__main__":
    app.run(debug=True)
